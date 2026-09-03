"""
OmniConverter PRO 4.0 Ultra - Core Python Conversion Engine
Supports high-fidelity conversion across Documents, Images, Audio, Video, Data Formats, and Code.
Integrates with File_Converter_Pro modules when present, with native fallback using Python standard libraries & PyPDF, Docx, OpenPyXL, Pandas, Pillow, etc.
"""

from __future__ import annotations
import os
import sys
import io
import time
import json
import base64
import re
import shutil
import zipfile
import tempfile
import subprocess
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

# Attempt to import File_Converter_Pro engine if available
try:
    sys.path.insert(0, str(Path(__file__).parent / "File_Converter_Pro"))
    from converter.converters import AdvancedConverterEngine
    FCP_ENGINE = AdvancedConverterEngine()
    HAS_FCP = True
except Exception as e:
    FCP_ENGINE = None
    HAS_FCP = False

# Third-party standard libraries check
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    from PIL import Image, ImageOps, ImageFilter
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

try:
    import pymupdf
    HAS_PYMUPDF = True
except ImportError:
    try:
        import fitz as pymupdf
        HAS_PYMUPDF = True
    except ImportError:
        pymupdf = None
        HAS_PYMUPDF = False

try:
    from rapidocr_onnxruntime import RapidOCR
    RAPID_OCR_INSTANCE = RapidOCR()
    HAS_RAPIDOCR = True
except Exception:
    RAPID_OCR_INSTANCE = None
    HAS_RAPIDOCR = False

try:
    import pytesseract
    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False


def check_ffmpeg() -> bool:
    """Checks whether FFmpeg executable is present on the system PATH."""
    return shutil.which("ffmpeg") is not None or shutil.which("ffmpeg.exe") is not None


class ConversionResult:
    def __init__(self, success: bool, output_path: str = "", filename: str = "", mime_type: str = "", error: str = "", elapsed: float = 0.0):
        self.success = success
        self.output_path = output_path
        self.filename = filename
        self.mime_type = mime_type
        self.error = error
        self.elapsed = elapsed

    def to_dict(self) -> Dict[str, Any]:
        size = os.path.getsize(self.output_path) if (self.output_path and os.path.exists(self.output_path)) else 0
        return {
            "success": self.success,
            "output_path": self.output_path,
            "filename": self.filename,
            "mime_type": self.mime_type,
            "error": self.error,
            "elapsed": round(self.elapsed, 3),
            "size": size
        }


class OmniConverterEngine:
    """Main Python conversion dispatcher for OmniConverter PRO 4.0."""

    MIME_MAP = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain",
        "html": "text/html",
        "md": "text/markdown",
        "json": "application/json",
        "csv": "text/csv",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "webp": "image/webp",
        "bmp": "image/bmp",
        "ico": "image/x-icon",
        "gif": "image/gif",
        "mp3": "audio/mpeg",
        "wav": "audio/wav",
        "ogg": "audio/ogg",
        "flac": "audio/flac",
        "mp4": "video/mp4",
        "webm": "video/webm",
        "zip": "application/zip"
    }

    FORMAT_CATEGORIES = {
        "image": {
            "name": "Image Formats",
            "extensions": ["png", "jpg", "jpeg", "webp", "bmp", "gif", "ico", "tiff"],
            "targets": ["png", "jpg", "jpeg", "webp", "bmp", "ico", "pdf"]
        },
        "document": {
            "name": "Documents & Text",
            "extensions": ["pdf", "docx", "txt", "html", "md"],
            "targets": ["pdf", "docx", "txt", "html", "md", "png", "jpg"]
        },
        "audio": {
            "name": "Audio",
            "extensions": ["mp3", "wav", "ogg", "flac", "aac", "m4a", "opus"],
            "targets": ["mp3", "wav", "ogg", "flac", "aac"]
        },
        "video": {
            "name": "Video",
            "extensions": ["mp4", "webm", "mkv", "avi", "mov", "flv"],
            "targets": ["mp4", "webm", "gif", "mp3", "wav", "aac", "flac", "ogg"]
        },
        "data": {
            "name": "Data & Spreadsheets",
            "extensions": ["csv", "json", "xlsx"],
            "targets": ["csv", "json", "xlsx", "txt"]
        }
    }

    def get_supported_formats(self) -> Dict[str, Any]:
        """Returns categories, supported extensions and target maps."""
        all_exts = {}
        for cat_id, cat_data in self.FORMAT_CATEGORIES.items():
            for ext in cat_data["extensions"]:
                all_exts[ext] = {
                    "category": cat_id,
                    "category_name": cat_data["name"],
                    "targets": cat_data["targets"]
                }
        return {
            "categories": self.FORMAT_CATEGORIES,
            "extensions": all_exts
        }

    def convert_file(self, input_path_or_bytes: str | bytes, input_filename: str, target_format: str, options: Optional[Dict[str, Any]] = None) -> ConversionResult:
        """Converts an input file to target format, streaming output to a temporary disk file to keep RAM usage low."""
        t0 = time.time()
        options = options or {}
        target_format = target_format.lower().lstrip('.')
        
        # Sanitize input filename to prevent path traversal attacks
        safe_filename = Path(input_filename).name
        input_ext = Path(safe_filename).suffix.lower().lstrip('.')

        temp_dir = tempfile.mkdtemp(prefix="omni_conv_")
        if isinstance(input_path_or_bytes, bytes):
            src_path = os.path.join(temp_dir, safe_filename)
            with open(src_path, "wb") as f:
                f.write(input_path_or_bytes)
        else:
            src_path = input_path_or_bytes

        out_stem = Path(safe_filename).stem
        out_filename = f"{out_stem}_converted.{target_format}"
        out_path = os.path.join(temp_dir, out_filename)

        try:
            # 1. Check FFmpeg dependency for audio/video media conversions
            if target_format in ["mp3", "wav", "ogg", "flac", "mp4", "webm"] and input_ext in ["mp4", "mkv", "avi", "mov", "webm", "wav", "mp3", "flac", "ogg"]:
                if not check_ffmpeg() and not HAS_FCP:
                    raise RuntimeError("FFmpeg executable not found in system PATH. Required for audio/video media conversion.")

            # 2. Try File Converter Pro Engine first if available
            if HAS_FCP and FCP_ENGINE is not None:
                conv_type = self._determine_fcp_type(input_ext, target_format)
                if conv_type:
                    try:
                        res = FCP_ENGINE.convert(conv_type, src_path, temp_dir)
                        if res and res.success and os.path.exists(res.target):
                            elapsed = time.time() - t0
                            return ConversionResult(
                                success=True,
                                output_path=res.target,
                                filename=Path(res.target).name,
                                mime_type=self.MIME_MAP.get(target_format, "application/octet-stream"),
                                elapsed=elapsed
                            )
                    except Exception as e_fcp:
                        print(f"[FCP Engine Fallback Warning]: {e_fcp}")

            # 3. Native Python Fallback Converters
            self._convert_native(src_path, out_path, input_ext, target_format, options, out_stem)
            
            elapsed = time.time() - t0
            return ConversionResult(
                success=True,
                output_path=out_path,
                filename=out_filename,
                mime_type=self.MIME_MAP.get(target_format, "application/octet-stream"),
                elapsed=elapsed
            )

        except Exception as e:
            elapsed = time.time() - t0
            shutil.rmtree(temp_dir, ignore_errors=True)
            return ConversionResult(
                success=False,
                error=str(e),
                elapsed=elapsed
            )

    def _determine_fcp_type(self, src_ext: str, dst_ext: str) -> Optional[str]:
        if src_ext == "pdf" and dst_ext in ["docx", "txt", "html"]:
            return f"pdf_to_{dst_ext}"
        if src_ext == "docx" and dst_ext in ["pdf", "txt", "html"]:
            return f"docx_to_{dst_ext}"
        if src_ext in ["png", "jpg", "jpeg", "webp", "bmp"] and dst_ext in ["png", "jpeg", "jpg", "webp", "bmp", "ico", "pdf"]:
            return f"image_to_{dst_ext}"
        if src_ext in ["csv", "json", "xlsx"] and dst_ext in ["csv", "json", "pdf", "xlsx"]:
            return f"{src_ext}_to_{dst_ext}"
        return None

    def _convert_native(self, src_path: str, out_path: str, src_ext: str, dst_ext: str, options: Dict[str, Any], stem: str):
        """Native python conversion handlers writing directly to out_path file."""

        # Same format identity copy
        if src_ext == dst_ext:
            shutil.copyfile(src_path, out_path)
            return

        # 1. IMAGE -> IMAGE & IMAGE -> PDF (Pillow)
        if src_ext in ["png", "jpg", "jpeg", "webp", "bmp", "gif", "ico", "tiff"]:
            if dst_ext == "pdf":
                if not HAS_PILLOW:
                    raise RuntimeError("Pillow module is required for image to PDF conversion.")
                with Image.open(src_path) as img:
                    if img.mode in ("RGBA", "P"):
                        img = img.convert("RGB")
                    img.save(out_path, format="PDF")
                return

            if dst_ext in ["png", "jpg", "jpeg", "webp", "bmp", "ico"]:
                if not HAS_PILLOW:
                    raise RuntimeError("Pillow module is required for image conversion.")
                
                with Image.open(src_path) as img:
                    if options.get("resize_width") and options.get("resize_height"):
                        img = img.resize((int(options["resize_width"]), int(options["resize_height"])), Image.Resampling.LANCZOS)
                    elif options.get("scale"):
                        scale = float(options["scale"])
                        img = img.resize((int(img.width * scale), int(img.height * scale)), Image.Resampling.LANCZOS)
                    
                    if options.get("rotate"):
                        img = img.rotate(float(options["rotate"]), expand=True)

                    if options.get("grayscale"):
                        img = img.convert("L")

                    target_fmt = "JPEG" if dst_ext in ["jpg", "jpeg"] else dst_ext.upper()
                    if target_fmt == "ICO":
                        img.save(out_path, format="ICO", sizes=[(32, 32), (64, 64), (128, 128)])
                    elif target_fmt == "PNG":
                        img.save(out_path, format="PNG", optimize=True)
                    else:
                        if img.mode in ("RGBA", "P") and target_fmt == "JPEG":
                            img = img.convert("RGB")
                        quality = int(options.get("quality", options.get("compressionQuality", 90)))
                        img.save(out_path, format=target_fmt, quality=quality)
                return

        # 2. PDF -> IMAGES (PyMuPDF)
        if src_ext == "pdf" and dst_ext in ["png", "jpg", "jpeg", "webp"]:
            if not HAS_PYMUPDF:
                raise RuntimeError("PyMuPDF is required for PDF to image conversion.")
            doc = pymupdf.open(src_path)
            dpi = int(options.get("dpi", 150))
            if len(doc) == 0:
                doc.close()
                raise ValueError("PDF document has no pages.")
            elif len(doc) == 1:
                pix = doc[0].get_pixmap(dpi=dpi)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                fmt = "JPEG" if dst_ext in ["jpg", "jpeg"] else dst_ext.upper()
                img.save(out_path, format=fmt)
            else:
                zip_temp = f"{out_path}.zip"
                with zipfile.ZipFile(zip_temp, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, page in enumerate(doc):
                        pix = page.get_pixmap(dpi=dpi)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        img_buf = io.BytesIO()
                        fmt = "JPEG" if dst_ext in ["jpg", "jpeg"] else dst_ext.upper()
                        img.save(img_buf, format=fmt)
                        zf.writestr(f"{stem}_page_{i+1}.{dst_ext}", img_buf.getvalue())
                shutil.move(zip_temp, out_path)
            doc.close()
            return

        # 3. TEXT & MARKDOWN -> PDF (PyMuPDF)
        if src_ext in ["txt", "md"] and dst_ext == "pdf":
            if HAS_PYMUPDF:
                with open(src_path, "r", encoding="utf-8", errors="ignore") as tf:
                    content = tf.read()
                doc = pymupdf.open()
                lines = content.splitlines() or [""]
                lines_per_page = 42
                for p_start in range(0, max(1, len(lines)), lines_per_page):
                    page = doc.new_page(width=595, height=842)
                    page_text = "\n".join(lines[p_start:p_start + lines_per_page])
                    page.insert_text((50, 72), page_text, fontsize=10)
                doc.save(out_path)
                doc.close()
                return

        # 4. MARKDOWN -> HTML
        if src_ext == "md" and dst_ext == "html":
            with open(src_path, "r", encoding="utf-8", errors="ignore") as rf:
                md_lines = rf.read().splitlines()
            html_lines = [
                f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{stem}</title>",
                "<style>body{font-family:system-ui,sans-serif;line-height:1.6;padding:2rem;max-width:800px;margin:auto;color:#1e293b;}</style></head><body>"
            ]
            for l in md_lines:
                if l.startswith("# "):
                    html_lines.append(f"<h1>{l[2:]}</h1>")
                elif l.startswith("## "):
                    html_lines.append(f"<h2>{l[3:]}</h2>")
                elif l.startswith("### "):
                    html_lines.append(f"<h3>{l[4:]}</h3>")
                elif l.strip():
                    html_lines.append(f"<p>{l}</p>")
            html_lines.append("</body></html>")
            with open(out_path, "w", encoding="utf-8") as wf:
                wf.write("\n".join(html_lines))
            return

        # 5. AUDIO & VIDEO CONVERSIONS (FFmpeg)
        if (src_ext in ["mp4", "mkv", "avi", "mov", "webm", "mp3", "wav", "aac", "flac", "ogg"] and 
            dst_ext in ["mp4", "webm", "gif", "mp3", "wav", "aac", "flac", "ogg"]):
            if check_ffmpeg():
                cmd = ["ffmpeg", "-y", "-i", src_path]
                
                # Video Options
                if dst_ext in ["mp4", "webm", "gif"]:
                    vq = str(options.get("video_quality") or options.get("videoQuality") or "Original")
                    if "1080" in vq:
                        cmd.extend(["-vf", "scale=-2:1080"])
                    elif "720" in vq:
                        cmd.extend(["-vf", "scale=-2:720"])
                    elif "480" in vq:
                        cmd.extend(["-vf", "scale=-2:480"])
                    elif "360" in vq:
                        cmd.extend(["-vf", "scale=-2:360"])
                    
                    if options.get("strip_audio") or options.get("stripAudio"):
                        cmd.append("-an")
                
                # Audio / Audio Extraction Options
                if dst_ext in ["mp3", "wav", "aac", "flac", "ogg"]:
                    cmd.append("-vn")
                    ab = str(options.get("audio_bitrate") or options.get("audioBitrate") or "320k").replace("kbps", "").replace("k", "")
                    cmd.extend(["-b:a", f"{ab}k"])
                
                cmd.append(out_path)
                res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                if res.returncode == 0 and os.path.exists(out_path):
                    return
                else:
                    err = res.stderr.decode("utf-8", errors="ignore")
                    raise RuntimeError(f"FFmpeg conversion failed: {err[:200]}")

        # 6. PDF -> TXT (PyMuPDF / RapidOCR / PyPDF)
        if src_ext == "pdf" and dst_ext == "txt":
            ocr_res = self.ocr_pdf(src_path, page_range="all", dpi=150, force_ocr=options.get("force_ocr", False))
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(ocr_res.get("text", ""))
            return

        # 7. PDF -> HTML (PyPDF)
        if src_ext == "pdf" and dst_ext == "html":
            if not HAS_PYPDF:
                raise RuntimeError("pypdf library is required for PDF to HTML conversion.")
            reader = pypdf.PdfReader(src_path)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write("<!DOCTYPE html><html><head><meta charset='utf-8'><title>PDF Export</title>")
                f.write("<style>body{font-family:sans-serif;padding:2rem;background:#0f172a;color:#f8fafc;}")
                f.write(".page{background:#1e293b;margin-bottom:1.5rem;padding:1.5rem;border-radius:1rem;border:1px solid #334155;}</style></head><body>")
                for i, page in enumerate(reader.pages):
                    txt = (page.extract_text() or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    f.write(f"<div class='page'><h3>Page {i+1}</h3><pre>{txt}</pre></div>")
                f.write("</body></html>")
            return

        # 8. DOCX -> TXT (python-docx)
        if src_ext == "docx" and dst_ext == "txt":
            if not HAS_DOCX:
                raise RuntimeError("python-docx library is required for DOCX conversion.")
            doc = docx.Document(src_path)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write("\n\n".join([p.text for p in doc.paragraphs if p.text]))
            return

        # 9. DOCX -> HTML (python-docx)
        if src_ext == "docx" and dst_ext == "html":
            if not HAS_DOCX:
                raise RuntimeError("python-docx library is required for DOCX to HTML conversion.")
            doc = docx.Document(src_path)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write("<!DOCTYPE html><html><body>")
                for p in doc.paragraphs:
                    if p.text:
                        f.write(f"<p>{p.text}</p>")
                f.write("</body></html>")
            return

        # 10. DATA CONVERSIONS (CSV, JSON, XLSX via Pandas)
        if src_ext == "csv" and dst_ext == "json":
            if HAS_PANDAS:
                df = pd.read_csv(src_path)
                df.to_json(out_path, orient="records", indent=2)
            else:
                with open(src_path, "r", encoding="utf-8") as rf:
                    import csv
                    reader = list(csv.DictReader(rf))
                with open(out_path, "w", encoding="utf-8") as wf:
                    json.dump(reader, wf, indent=2)
            return

        if src_ext == "csv" and dst_ext == "xlsx":
            if HAS_PANDAS:
                df = pd.read_csv(src_path)
                df.to_excel(out_path, index=False)
                return
            elif HAS_OPENPYXL:
                import csv
                wb = openpyxl.Workbook()
                ws = wb.active
                with open(src_path, "r", encoding="utf-8") as cf:
                    for row in csv.reader(cf):
                        ws.append(row)
                wb.save(out_path)
                return

        if src_ext == "json" and dst_ext == "csv":
            with open(src_path, "r", encoding="utf-8") as rf:
                data = json.load(rf)
            if isinstance(data, list) and len(data) > 0 and HAS_PANDAS:
                df = pd.DataFrame(data)
                df.to_csv(out_path, index=False)
            elif isinstance(data, list) and len(data) > 0:
                import csv
                with open(out_path, "w", encoding="utf-8", newline="") as wf:
                    writer = csv.DictWriter(wf, fieldnames=data[0].keys())
                    writer.writeheader()
                    writer.writerows(data)
            return

        if src_ext == "xlsx" and dst_ext in ["csv", "json"]:
            if HAS_PANDAS:
                df = pd.read_excel(src_path)
                if dst_ext == "csv":
                    df.to_csv(out_path, index=False)
                else:
                    df.to_json(out_path, orient="records", indent=2)
            return

        # 11. TXT <-> JSON
        if src_ext in ["txt", "text"] and dst_ext == "json":
            with open(src_path, "r", encoding="utf-8", errors="ignore") as tf:
                content = tf.read()
            try:
                parsed = json.loads(content)
                with open(out_path, "w", encoding="utf-8") as jf:
                    json.dump(parsed, jf, indent=2)
            except Exception:
                with open(out_path, "w", encoding="utf-8") as jf:
                    json.dump({"filename": Path(src_path).name, "content": content, "lines": content.splitlines()}, jf, indent=2)
            return

        if src_ext == "json" and dst_ext in ["txt", "text"]:
            with open(src_path, "r", encoding="utf-8", errors="ignore") as jf:
                data = json.load(jf)
            with open(out_path, "w", encoding="utf-8") as tf:
                if isinstance(data, dict) and "content" in data:
                    tf.write(str(data["content"]))
                else:
                    tf.write(json.dumps(data, indent=2))
            return

        # If completely unsupported, raise ValueError instead of creating corrupted copies
        raise ValueError(f"Direct conversion from '.{src_ext}' to '.{dst_ext}' is not supported.")

    # ==================== OMNI PDF ENGINE MODULES ====================
    
    @staticmethod
    def parse_page_ranges(range_str: str, total_pages: int) -> List[int]:
        """
        Parses page strings like '1-3, 5', 'odd', 'even', 'all' into zero-indexed page numbers.
        """
        if not range_str or range_str.strip().lower() == "all":
            return list(range(total_pages))
        
        range_str = range_str.strip().lower()
        if range_str == "odd":
            return [i for i in range(total_pages) if (i + 1) % 2 != 0]
        if range_str == "even":
            return [i for i in range(total_pages) if (i + 1) % 2 == 0]

        selected_pages = set()
        parts = range_str.split(",")
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if "-" in part:
                sub = part.split("-")
                if len(sub) == 2 and sub[0].isdigit() and sub[1].isdigit():
                    start = max(1, int(sub[0]))
                    end = min(total_pages, int(sub[1]))
                    for p in range(start, end + 1):
                        selected_pages.add(p - 1)
            elif part.isdigit():
                p = int(part)
                if 1 <= p <= total_pages:
                    selected_pages.add(p - 1)

        return sorted(list(selected_pages)) if selected_pages else list(range(total_pages))

    @staticmethod
    def _handle_encrypted_reader(reader: pypdf.PdfReader, password: str = ""):
        if reader.is_encrypted:
            if password:
                decrypted = reader.decrypt(password)
                if decrypted == 0:
                    raise ValueError("Incorrect password for protected PDF.")
            else:
                raise ValueError("PDF is password protected. Please provide a password.")

    def merge_pdfs(self, pdf_paths: List[str], out_path: str, passwords: Optional[List[str]] = None):
        if not HAS_PYPDF:
            raise RuntimeError("pypdf library is required for PDF Merge.")
        writer = pypdf.PdfWriter()
        passwords = passwords or []
        for i, path in enumerate(pdf_paths):
            pwd = passwords[i] if i < len(passwords) else ""
            reader = pypdf.PdfReader(path)
            self._handle_encrypted_reader(reader, pwd)
            writer.append(reader)
        with open(out_path, "wb") as f:
            writer.write(f)
        writer.close()

    def split_pdf(self, pdf_path: str, page_range: str, out_path: str, mode: str = "single_pdf", password: str = "") -> str:
        if not HAS_PYPDF:
            raise RuntimeError("pypdf library is required for PDF Split.")
        reader = pypdf.PdfReader(pdf_path)
        self._handle_encrypted_reader(reader, password)

        total_pages = len(reader.pages)
        indices = self.parse_page_ranges(page_range, total_pages)

        if mode == "zip":
            # Output a ZIP file containing individual PDF pages
            temp_dir = tempfile.mkdtemp(prefix="omni_split_zip_")
            stem = Path(pdf_path).stem
            zip_out_path = out_path if out_path.endswith(".zip") else f"{out_path}.zip"
            
            with zipfile.ZipFile(zip_out_path, "w", zipfile.ZIP_DEFLATED) as zf:
                for idx in indices:
                    w = pypdf.PdfWriter()
                    w.add_page(reader.pages[idx])
                    page_filename = f"{stem}_page_{idx+1}.pdf"
                    page_filepath = os.path.join(temp_dir, page_filename)
                    with open(page_filepath, "wb") as pf:
                        w.write(pf)
                    zf.write(page_filepath, arcname=page_filename)
            shutil.rmtree(temp_dir, ignore_errors=True)
            return zip_out_path
        else:
            # Output a single PDF with all selected pages
            writer = pypdf.PdfWriter()
            for idx in indices:
                writer.add_page(reader.pages[idx])
            with open(out_path, "wb") as f:
                writer.write(f)
            return out_path

    def compress_pdf(self, pdf_path: str, out_path: str, level: str = "medium", password: str = ""):
        if not HAS_PYPDF:
            raise RuntimeError("pypdf library is required for PDF Compress.")
        reader = pypdf.PdfReader(pdf_path)
        self._handle_encrypted_reader(reader, password)

        writer = pypdf.PdfWriter()
        scale_factor = 0.8 if level == "low" else (0.5 if level == "medium" else 0.3)
        quality = 80 if level == "low" else (60 if level == "medium" else 45)

        for page in reader.pages:
            page.compress_content_streams()
            # If Pillow is present and aggressive compression requested, downsample images
            if HAS_PILLOW and level in ["medium", "high"]:
                try:
                    for count, image_file_object in enumerate(page.images):
                        img = Image.open(io.BytesIO(image_file_object.data))
                        new_size = (max(1, int(img.width * scale_factor)), max(1, int(img.height * scale_factor)))
                        img = img.resize(new_size, Image.Resampling.LANCZOS)
                        if img.mode in ("RGBA", "P"):
                            img = img.convert("RGB")
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format="JPEG", quality=quality)
                        page.images[count].replace(img, img_byte_arr.getvalue())
                except Exception as e_img:
                    print(f"[PDF Compress Image Downsample Warning]: {e_img}")
            writer.add_page(page)

        with open(out_path, "wb") as f:
            writer.write(f)

    def protect_pdf(self, pdf_path: str, password: str, out_path: str):
        if not HAS_PYPDF:
            raise RuntimeError("pypdf library is required for PDF Protect.")
        reader = pypdf.PdfReader(pdf_path)
        writer = pypdf.PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(password)
        with open(out_path, "wb") as f:
            writer.write(f)

    def unlock_pdf(self, pdf_path: str, password: str, out_path: str):
        if not HAS_PYPDF:
            raise RuntimeError("pypdf library is required for PDF Unlock.")
        reader = pypdf.PdfReader(pdf_path)
        self._handle_encrypted_reader(reader, password)
        writer = pypdf.PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with open(out_path, "wb") as f:
            writer.write(f)

    def rotate_pdf(self, pdf_path: str, angle: int, out_path: str, page_range: str = "all", password: str = ""):
        if not HAS_PYPDF:
            raise RuntimeError("pypdf library is required for PDF Rotate.")
        reader = pypdf.PdfReader(pdf_path)
        self._handle_encrypted_reader(reader, password)

        total_pages = len(reader.pages)
        rotate_indices = set(self.parse_page_ranges(page_range, total_pages))

        writer = pypdf.PdfWriter()
        for i, page in enumerate(reader.pages):
            if i in rotate_indices:
                page.rotate(angle)
            writer.add_page(page)
        with open(out_path, "wb") as f:
            writer.write(f)

    def extract_ocr_from_image(self, image_input: str | bytes | Image.Image) -> Dict[str, Any]:
        """Runs OCR on an image (path, bytes, or PIL Image) and returns text and line-level coordinates."""
        if isinstance(image_input, (str, Path)):
            img = Image.open(image_input)
        elif isinstance(image_input, bytes):
            img = Image.open(io.BytesIO(image_input))
        else:
            img = image_input

        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        lines = []
        full_text = ""
        avg_confidence = 1.0

        if HAS_RAPIDOCR and RAPID_OCR_INSTANCE is not None:
            import numpy as np
            img_np = np.array(img)
            ocr_res, _ = RAPID_OCR_INSTANCE(img_np)
            if ocr_res:
                confidences = []
                for item in ocr_res:
                    text_str = str(item[1]).strip()
                    try:
                        score = float(item[2])
                    except Exception:
                        score = 0.9
                    if text_str:
                        lines.append({
                            "text": text_str,
                            "box": item[0] if isinstance(item[0], list) else [],
                            "confidence": round(score, 3)
                        })
                        confidences.append(score)
                if confidences:
                    avg_confidence = round(sum(confidences) / len(confidences), 3)
                full_text = "\n".join(l["text"] for l in lines)
        elif HAS_PYTESSERACT:
            full_text = pytesseract.image_to_string(img)
            lines = [{"text": l, "confidence": 0.85} for l in full_text.splitlines() if l.strip()]

        return {
            "text": full_text,
            "lines": lines,
            "confidence": avg_confidence,
            "line_count": len(lines)
        }

    def ocr_pdf(self, pdf_path_or_bytes: str | bytes, page_range: str = "all", dpi: int = 150, force_ocr: bool = False, password: str = "") -> Dict[str, Any]:
        """Extracts text from PDF pages using PyMuPDF and RapidOCR. Falls back to PyPDF."""
        t0 = time.time()
        pages_data = []
        all_text_list = []

        if isinstance(pdf_path_or_bytes, bytes):
            doc = pymupdf.open(stream=pdf_path_or_bytes, filetype="pdf") if HAS_PYMUPDF else None
            pypdf_bytes = io.BytesIO(pdf_path_or_bytes)
        else:
            doc = pymupdf.open(pdf_path_or_bytes) if HAS_PYMUPDF else None
            pypdf_bytes = pdf_path_or_bytes

        if doc is not None and password:
            doc.authenticate(password)

        total_pages = len(doc) if doc else (len(pypdf.PdfReader(pypdf_bytes).pages) if HAS_PYPDF else 0)
        target_indices = self.parse_page_ranges(page_range, total_pages) if total_pages > 0 else []

        overall_confidences = []

        if HAS_PYMUPDF and doc is not None:
            for idx in target_indices:
                page = doc[idx]
                digital_text = page.get_text().strip()

                # If digital text exists and force_ocr is False, use high-speed digital text
                if digital_text and not force_ocr and len(digital_text) > 15:
                    page_text = digital_text
                    confidence = 1.0
                    method = "digital"
                    lines = [{"text": l, "confidence": 1.0} for l in page_text.splitlines() if l.strip()]
                else:
                    # Render raster image from PDF page for OCR
                    pix = page.get_pixmap(dpi=dpi)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_res = self.extract_ocr_from_image(img)
                    page_text = ocr_res["text"] or digital_text
                    confidence = ocr_res["confidence"]
                    method = "ocr" if ocr_res["text"] else "digital_fallback"
                    lines = ocr_res["lines"]

                overall_confidences.append(confidence)
                pages_data.append({
                    "page_number": idx + 1,
                    "text": page_text,
                    "confidence": confidence,
                    "method": method,
                    "line_count": len(lines),
                    "lines": lines
                })
                all_text_list.append(f"--- PAGE {idx + 1} ---\n{page_text}")

            doc.close()
        elif HAS_PYPDF:
            reader = pypdf.PdfReader(pypdf_bytes)
            self._handle_encrypted_reader(reader, password)
            for idx in target_indices:
                page = reader.pages[idx]
                page_text = page.extract_text() or ""
                method = "digital"
                confidence = 1.0

                if not page_text and len(page.images) > 0:
                    img_texts = []
                    for img_obj in page.images:
                        res = self.extract_ocr_from_image(img_obj.data)
                        if res["text"]:
                            img_texts.append(res["text"])
                    if img_texts:
                        page_text = "\n".join(img_texts)
                        method = "pypdf_ocr"
                        confidence = 0.9

                pages_data.append({
                    "page_number": idx + 1,
                    "text": page_text,
                    "confidence": confidence,
                    "method": method,
                    "line_count": len(page_text.splitlines()),
                    "lines": [{"text": l, "confidence": confidence} for l in page_text.splitlines() if l.strip()]
                })
                all_text_list.append(f"--- PAGE {idx + 1} ---\n{page_text}")

        elapsed = time.time() - t0
        avg_conf = round(sum(overall_confidences) / max(1, len(overall_confidences)), 3) if overall_confidences else 1.0
        combined_text = "\n\n".join(all_text_list)

        return {
            "success": True,
            "total_pages": total_pages,
            "processed_pages": len(pages_data),
            "text": combined_text,
            "confidence": avg_conf,
            "pages": pages_data,
            "elapsed": round(elapsed, 3)
        }

    def ocr_document(self, file_path_or_bytes: str | bytes, filename: str, page_range: str = "all", force_ocr: bool = False, password: str = "") -> Dict[str, Any]:
        """Unified OCR dispatcher for PDF documents and image files."""
        ext = Path(filename).suffix.lower().lstrip(".")
        if ext == "pdf":
            return self.ocr_pdf(file_path_or_bytes, page_range=page_range, force_ocr=force_ocr, password=password)
        elif ext in ["png", "jpg", "jpeg", "webp", "bmp", "tiff", "tif"]:
            t0 = time.time()
            ocr_res = self.extract_ocr_from_image(file_path_or_bytes)
            elapsed = time.time() - t0
            return {
                "success": True,
                "total_pages": 1,
                "processed_pages": 1,
                "text": ocr_res["text"],
                "confidence": ocr_res["confidence"],
                "pages": [{
                    "page_number": 1,
                    "text": ocr_res["text"],
                    "confidence": ocr_res["confidence"],
                    "method": "image_ocr",
                    "line_count": ocr_res["line_count"],
                    "lines": ocr_res["lines"]
                }],
                "elapsed": round(elapsed, 3)
            }
    def pdf_merge(self, files: List[Tuple[str, bytes] | str], out_path: Optional[str] = None, passwords: Optional[List[str]] = None) -> str:
        """High-level programmatic wrapper to merge multiple PDF files."""
        temp_dir = tempfile.mkdtemp(prefix="omni_merge_")
        dest = out_path or os.path.join(temp_dir, "OmniConverter_Merged.pdf")
        paths = []
        for item in files:
            if isinstance(item, str):
                paths.append(item)
            elif isinstance(item, tuple):
                p = os.path.join(temp_dir, item[0])
                with open(p, "wb") as f:
                    f.write(item[1])
                paths.append(p)
        self.merge_pdfs(paths, dest, passwords)
        return dest

    def pdf_split(self, file_path_or_bytes: str | bytes, page_range: str = "all", mode: str = "single_pdf", out_path: Optional[str] = None, password: str = "") -> str:
        """High-level programmatic wrapper to split PDF pages."""
        temp_dir = tempfile.mkdtemp(prefix="omni_split_")
        if isinstance(file_path_or_bytes, bytes):
            src_path = os.path.join(temp_dir, "input.pdf")
            with open(src_path, "wb") as f:
                f.write(file_path_or_bytes)
        else:
            src_path = str(file_path_or_bytes)
        ext = ".zip" if mode == "zip" else ".pdf"
        dest = out_path or os.path.join(temp_dir, f"split_output{ext}")
        return self.split_pdf(src_path, page_range, dest, mode, password)

    def pdf_compress(self, file_path_or_bytes: str | bytes, level: str = "medium", out_path: Optional[str] = None, password: str = "") -> str:
        """High-level programmatic wrapper to compress PDF size."""
        temp_dir = tempfile.mkdtemp(prefix="omni_compress_")
        if isinstance(file_path_or_bytes, bytes):
            src_path = os.path.join(temp_dir, "input.pdf")
            with open(src_path, "wb") as f:
                f.write(file_path_or_bytes)
        else:
            src_path = str(file_path_or_bytes)
        dest = out_path or os.path.join(temp_dir, "compressed_output.pdf")
        self.compress_pdf(src_path, dest, level, password)
        return dest

    def pdf_protect(self, file_path_or_bytes: str | bytes, password: str, out_path: Optional[str] = None) -> str:
        """High-level programmatic wrapper to encrypt PDF with password."""
        temp_dir = tempfile.mkdtemp(prefix="omni_protect_")
        if isinstance(file_path_or_bytes, bytes):
            src_path = os.path.join(temp_dir, "input.pdf")
            with open(src_path, "wb") as f:
                f.write(file_path_or_bytes)
        else:
            src_path = str(file_path_or_bytes)
        dest = out_path or os.path.join(temp_dir, "protected_output.pdf")
        self.protect_pdf(src_path, dest, password)
        return dest

    def pdf_unlock(self, file_path_or_bytes: str | bytes, password: str, out_path: Optional[str] = None) -> str:
        """High-level programmatic wrapper to decrypt protected PDF."""
        temp_dir = tempfile.mkdtemp(prefix="omni_unlock_")
        if isinstance(file_path_or_bytes, bytes):
            src_path = os.path.join(temp_dir, "input.pdf")
            with open(src_path, "wb") as f:
                f.write(file_path_or_bytes)
        else:
            src_path = str(file_path_or_bytes)
        dest = out_path or os.path.join(temp_dir, "unlocked_output.pdf")
        self.unlock_pdf(src_path, dest, password)
        return dest

    def pdf_rotate(self, file_path_or_bytes: str | bytes, angle: int = 90, page_range: str = "all", out_path: Optional[str] = None, password: str = "") -> str:
        """High-level programmatic wrapper to rotate PDF pages."""
        temp_dir = tempfile.mkdtemp(prefix="omni_rotate_")
        if isinstance(file_path_or_bytes, bytes):
            src_path = os.path.join(temp_dir, "input.pdf")
            with open(src_path, "wb") as f:
                f.write(file_path_or_bytes)
        else:
            src_path = str(file_path_or_bytes)
        dest = out_path or os.path.join(temp_dir, "rotated_output.pdf")
        self.rotate_pdf(src_path, dest, angle, page_range, password)
        return dest

    def convert_batch(self, files: List[Tuple[str, bytes]], target_format: str, options: Optional[Dict[str, Any]] = None) -> str:
        """Converts multiple files in batch and returns the file path of a ZIP archive."""
        temp_dir = tempfile.mkdtemp(prefix="omni_batch_")
        zip_path = os.path.join(temp_dir, f"OmniConverter_Batch_{target_format.upper()}.zip")

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, content in files:
                res = self.convert_file(content, fname, target_format, options)
                if res.success and os.path.exists(res.output_path):
                    zf.write(res.output_path, arcname=res.filename)
                else:
                    err_msg = f"Failed to convert {fname}: {res.error}"
                    zf.writestr(f"ERROR_{fname}.txt", err_msg.encode("utf-8"))
        return zip_path


# Global engine instance
converter_engine = OmniConverterEngine()


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        prog="omniconverter",
        description="OmniConverter CLI: Universal File Converter, PDF Suite & Neural OCR"
    )
    parser.add_argument("files", nargs="+", help="Input file path(s)")
    parser.add_argument("-t", "--target", help="Target conversion format (e.g. pdf, png, webp, mp3, docx)")
    parser.add_argument("-o", "--output", help="Destination file or directory path")
    parser.add_argument("--merge", action="store_true", help="Merge multiple PDF inputs into a single document")
    parser.add_argument("--split", metavar="RANGE", help="Split PDF pages (e.g. 1-3, odd)")
    parser.add_argument("--compress", choices=["low", "medium", "high"], help="Compress PDF size")
    parser.add_argument("--protect", metavar="PASSWORD", help="Encrypt PDF with password")
    parser.add_argument("--unlock", metavar="PASSWORD", help="Decrypt password-protected PDF")
    parser.add_argument("--rotate", type=int, choices=[90, 180, 270], help="Rotate PDF pages clockwise")
    parser.add_argument("--ocr", action="store_true", help="Extract text via RapidOCR neural engine")

    args = parser.parse_args()

    try:
        # 1. MERGE PDFs
        if args.merge:
            pdf_data = [(Path(f).name, open(f, "rb").read()) for f in args.files]
            out_path = converter_engine.pdf_merge(pdf_data)
            dest = args.output or "merged_output.pdf"
            shutil.copyfile(out_path, dest)
            print(f"[SUCCESS] Merged {len(pdf_data)} PDFs -> {dest}")

        # 2. SPLIT PDF
        elif args.split:
            fpath = args.files[0]
            out_path = converter_engine.pdf_split(open(fpath, "rb").read(), page_range=args.split)
            dest = args.output or f"split_{Path(fpath).name}"
            shutil.copyfile(out_path, dest)
            print(f"[SUCCESS] Split PDF ({args.split}) -> {dest}")

        # 3. COMPRESS PDF
        elif args.compress:
            fpath = args.files[0]
            out_path = converter_engine.pdf_compress(open(fpath, "rb").read(), level=args.compress)
            dest = args.output or f"compressed_{Path(fpath).name}"
            shutil.copyfile(out_path, dest)
            print(f"[SUCCESS] Compressed PDF ({args.compress}) -> {dest}")

        # 4. PROTECT PDF
        elif args.protect:
            fpath = args.files[0]
            out_path = converter_engine.pdf_protect(open(fpath, "rb").read(), password=args.protect)
            dest = args.output or f"protected_{Path(fpath).name}"
            shutil.copyfile(out_path, dest)
            print(f"[SUCCESS] Protected PDF -> {dest}")

        # 5. UNLOCK PDF
        elif args.unlock:
            fpath = args.files[0]
            out_path = converter_engine.pdf_unlock(open(fpath, "rb").read(), password=args.unlock)
            dest = args.output or f"unlocked_{Path(fpath).name}"
            shutil.copyfile(out_path, dest)
            print(f"[SUCCESS] Unlocked PDF -> {dest}")

        # 6. ROTATE PDF
        elif args.rotate:
            fpath = args.files[0]
            out_path = converter_engine.pdf_rotate(open(fpath, "rb").read(), angle=args.rotate)
            dest = args.output or f"rotated_{Path(fpath).name}"
            shutil.copyfile(out_path, dest)
            print(f"[SUCCESS] Rotated PDF ({args.rotate} deg) -> {dest}")

        # 7. OCR TEXT EXTRACTION
        elif args.ocr:
            fpath = args.files[0]
            ocr_res = converter_engine.ocr_document(open(fpath, "rb").read(), Path(fpath).name)
            if args.output:
                with open(args.output, "w", encoding="utf-8") as out_f:
                    out_f.write(ocr_res["text"])
                print(f"[SUCCESS] OCR text extracted to {args.output} ({ocr_res['confidence'] * 100:.1f}% confidence)")
            else:
                print(ocr_res["text"])

        # 8. STANDARD CONVERSION
        elif args.target:
            for fpath in args.files:
                fname = Path(fpath).name
                res = converter_engine.convert_file(open(fpath, "rb").read(), fname, args.target)
                if res.success:
                    dest = args.output or res.filename
                    if os.path.isdir(dest):
                        dest = os.path.join(dest, res.filename)
                    shutil.copyfile(res.output_path, dest)
                    print(f"[SUCCESS] Converted {fname} -> {dest} in {res.elapsed}s")
                else:
                    print(f"[ERROR] Failed to convert {fname}: {res.error}", file=sys.stderr)

        else:
            parser.print_help()

    except Exception as e:
        print(f"[ERROR] {str(e)}", file=sys.stderr)
        sys.exit(1)



